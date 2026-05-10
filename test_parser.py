"""
test_parser.py — Smoke tests for sniffer_parser.py

Generates test DOCX fixtures, runs them through the parser,
and verifies hard errors, soft warnings, and structural output.

Usage:
    python test_parser.py

Test fixtures are preserved in testing/parser_fixtures/ for manual review.
"""
from __future__ import annotations

import copy
import json
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.enum.style import WD_STYLE_TYPE

sys.path.insert(0, str(Path(__file__).parent))
from sniffer_parser import parse_docx, to_dict

# ============================================================
# Infra
# ============================================================

passed = 0
failed = 0

FIXTURE_DIR = Path(__file__).parent / "testing" / "parser_fixtures"
TEMPLATE_PATH = Path("/mnt/user-data/outputs/branch-b-template-band4.docx")


def report(name: str, ok: bool, detail: str = ""):
    global passed, failed
    if ok:
        passed += 1
        print(f"  ✓ {name}")
    else:
        failed += 1
        print(f"  ✗ {name}")
        if detail:
            print(f"    → {detail}")


def save_fixture(doc, name: str) -> Path:
    """Save a fixture DOCX and return the path."""
    FIXTURE_DIR.mkdir(parents=True, exist_ok=True)
    path = FIXTURE_DIR / f"{name}.docx"
    doc.save(str(path))
    return path


def load_template() -> Document:
    return Document(str(TEMPLATE_PATH))


# ============================================================
# Helpers for modifying documents
# ============================================================

def find_para_by_text(doc, substring: str) -> Paragraph | None:
    for para in doc.paragraphs:
        if substring in para.text:
            return para
    return None


def find_para_by_style(doc, style_id: str, index: int = 0) -> Paragraph | None:
    count = 0
    for para in doc.paragraphs:
        sid = para.style.style_id if para.style else None
        if sid == style_id:
            if count == index:
                return para
            count += 1
    return None


def change_style(para, new_style_name: str):
    pPr = para._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        para._element.insert(0, pPr)
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        pStyle = OxmlElement("w:pStyle")
        pPr.insert(0, pStyle)
    pStyle.set(qn("w:val"), new_style_name)


def delete_paragraph(para):
    para._element.getparent().remove(para._element)


def add_tracked_change(doc):
    para = find_para_by_text(doc, "Hello")
    if para is None:
        return
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), "999")
    ins.set(qn("w:author"), "Test")
    ins.set(qn("w:date"), "2026-01-01T00:00:00Z")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "TRACKED"
    r.append(t)
    ins.append(r)
    para._element.append(ins)


def add_comment(doc):
    para = find_para_by_text(doc, "Hello")
    if para is None:
        return
    cs = OxmlElement("w:commentRangeStart")
    cs.set(qn("w:id"), "0")
    para._element.insert(0, cs)


# ============================================================
# Tests 01-03: Style violations (should hard-fail)
# ============================================================

def test_01_body_to_normal():
    """EL_Body changed to Normal — parser rejects."""
    doc = load_template()
    para = find_para_by_text(doc, "Hello")
    change_style(para, "Normal")
    path = save_fixture(doc, "01_body_to_normal")
    r = parse_docx(str(path))
    has_error = any("Non-EL_*" in e or "Normal" in e for e in r.hard_errors)
    report("01: EL_Body → Normal rejects", has_error,
           f"Errors: {r.hard_errors[:2]}" if not has_error else "")


def test_02_bullet_to_list_paragraph():
    """EL_Bullet changed to ListParagraph — parser rejects."""
    doc = load_template()
    para = find_para_by_text(doc, "Worker on leave")
    change_style(para, "ListParagraph")
    path = save_fixture(doc, "02_bullet_to_listparagraph")
    r = parse_docx(str(path))
    has_error = any("Non-EL_*" in e or "ListParagraph" in e or "List Paragraph" in e
                    for e in r.hard_errors)
    report("02: EL_Bullet → ListParagraph rejects", has_error,
           f"Errors: {r.hard_errors[:2]}" if not has_error else "")


def test_03_delete_el_rule():
    """Delete an EL_Rule — table loses sandwich, parser rejects."""
    doc = load_template()
    rule_para = find_para_by_style(doc, "EL_Rule", index=0)
    if rule_para is None:
        report("03: Delete EL_Rule rejects", False, "Could not find EL_Rule paragraph")
        return
    delete_paragraph(rule_para)
    path = save_fixture(doc, "03_delete_el_rule")
    r = parse_docx(str(path))
    has_error = any("EL_Rule" in e or "sandwich" in e or "Expected" in e
                    for e in r.hard_errors)
    report("03: Delete EL_Rule rejects", has_error,
           f"Errors: {r.hard_errors[:3]}" if not has_error else "")


# ============================================================
# Tests 04-05: Token violations (should hard-fail)
# ============================================================

def test_04_token_typo():
    """Typo in token [DEVICE_MODLE] — parser rejects as undefined."""
    doc = load_template()
    para = find_para_by_text(doc, "[DEVICE_MODEL]")
    if para is None:
        report("04: Token typo rejects", False, "Could not find [DEVICE_MODEL] paragraph")
        return
    for run in para._element.findall(f".//{qn('w:t')}"):
        if "[DEVICE_MODEL]" in run.text:
            run.text = run.text.replace("[DEVICE_MODEL]", "[DEVICE_MODLE]")
            break
    path = save_fixture(doc, "04_token_typo")
    r = parse_docx(str(path))
    has_error = any("DEVICE_MODLE" in e for e in r.hard_errors)
    report("04: Token typo [DEVICE_MODLE] rejects", has_error,
           f"Errors: {[e for e in r.hard_errors if 'MODLE' in e]}" if not has_error else "")


def test_05_override_new_token():
    """Per-language override defines token absent from global — parser rejects."""
    doc = load_template()
    body = doc.element.body
    children = list(body)

    for idx, child in enumerate(children):
        if child.tag == qn("w:p"):
            p = Paragraph(child, doc)
            if p.style and p.style.style_id == "EL_LegendHeading" and "Chinese" in p.text:
                for j in range(idx + 1, len(children)):
                    if children[j].tag == qn("w:tbl"):
                        tbl = children[j]
                        rows = tbl.findall(qn("w:tr"))
                        new_row = copy.deepcopy(rows[-1])
                        cells = new_row.findall(f".//{qn('w:t')}")
                        if cells:
                            cells[0].text = "<Bogus Token>"
                        tbl.append(new_row)
                        break
                break

    path = save_fixture(doc, "05_override_new_token")
    r = parse_docx(str(path))
    has_error = any("Bogus Token" in e and "not in global" in e for e in r.hard_errors)
    report("05: Override new token rejects", has_error,
           f"Errors: {[e for e in r.hard_errors if 'Bogus' in e]}" if not has_error else "")


# ============================================================
# Tests 06-07: Track Changes / Comments (should hard-fail)
# ============================================================

def test_06_track_changes():
    """Document with tracked changes — parser rejects."""
    doc = load_template()
    add_tracked_change(doc)
    path = save_fixture(doc, "06_track_changes")
    r = parse_docx(str(path))
    has_error = any("tracked changes" in e.lower() or "track changes" in e.lower()
                    for e in r.hard_errors)
    report("06: Track Changes rejects", has_error,
           f"Errors: {r.hard_errors[:2]}" if not has_error else "")


def test_07_comments():
    """Document with comments — parser rejects."""
    doc = load_template()
    add_comment(doc)
    path = save_fixture(doc, "07_comments")
    r = parse_docx(str(path))
    has_error = any("comment" in e.lower() for e in r.hard_errors)
    report("07: Comments rejects", has_error,
           f"Errors: {r.hard_errors[:2]}" if not has_error else "")


# ============================================================
# Tests 08-09: Soft warnings
# ============================================================

def test_08_unused_token():
    """Token defined in legend but never used in body — soft warning."""
    doc = load_template()
    for para in doc.paragraphs:
        for run_elem in para._element.findall(f".//{qn('w:t')}"):
            if "[DEVICE_SERIAL_NUMBER]" in run_elem.text:
                run_elem.text = run_elem.text.replace("[DEVICE_SERIAL_NUMBER]", "REMOVED")
    path = save_fixture(doc, "08_unused_token")
    r = parse_docx(str(path))
    has_warning = any("DEVICE_SERIAL_NUMBER" in w and "never used" in w
                      for w in r.soft_warnings)
    report("08: Unused token warns", has_warning,
           f"Warnings: {r.soft_warnings}" if not has_warning else "")


def test_09_noop_override():
    """Override row identical to global — soft warning."""
    doc = load_template()
    body = doc.element.body
    children = list(body)
    in_chinese = False
    for idx, child in enumerate(children):
        if child.tag == qn("w:p"):
            p = Paragraph(child, doc)
            if p.style and p.style.style_id == "EL_LegendHeading" and "Chinese" in p.text:
                in_chinese = True
                continue
            if in_chinese and p.style and p.style.style_id in ("EL_LegendHeading", "EL_BandHeading"):
                break
        if in_chinese and child.tag == qn("w:tbl"):
            rows = child.findall(qn("w:tr"))
            global_displays = [
                "device return process",
                "Request to Unblock Device from Quarantine",
                "ChatNow in Teams"
            ]
            for row_idx, row in enumerate(rows[1:], 0):
                cells = row.findall(qn("w:tc"))
                if len(cells) >= 2 and row_idx < len(global_displays):
                    for t in cells[1].findall(f".//{qn('w:t')}"):
                        t.text = global_displays[row_idx]
            break

    path = save_fixture(doc, "09_noop_override")
    r = parse_docx(str(path))
    has_warning = any("identical to global" in w or "no-op" in w for w in r.soft_warnings)
    report("09: No-op override warns", has_warning,
           f"Warnings: {r.soft_warnings}" if not has_warning else "")


# ============================================================
# Tests 10-12: Structure
# ============================================================

def test_10_missing_global_legend():
    """Delete global LEGEND heading — parser rejects."""
    doc = load_template()
    legend_heading = find_para_by_style(doc, "EL_LegendHeading", index=0)
    if legend_heading and "LEGEND" == legend_heading.text.strip():
        delete_paragraph(legend_heading)
    path = save_fixture(doc, "10_missing_legend")
    r = parse_docx(str(path))
    has_error = any("missing" in e.lower() and "legend" in e.lower()
                    for e in r.hard_errors)
    report("10: Missing global LEGEND rejects", has_error,
           f"Errors: {r.hard_errors[:3]}" if not has_error else "")


def test_11_band_before_legend():
    """EL_BandHeading before LEGEND — parser rejects on ordering."""
    doc = load_template()
    body = doc.element.body

    for child in list(body):
        if child.tag == qn("w:p"):
            p = Paragraph(child, doc)
            if p.style and p.style.style_id == "EL_BandHeading":
                body.remove(child)
                body.insert(0, child)
                break

    path = save_fixture(doc, "11_band_before_legend")
    r = parse_docx(str(path))
    has_error = any("before" in e.lower() and "legend" in e.lower()
                    for e in r.hard_errors)
    report("11: Band before LEGEND rejects on ordering", has_error,
           f"Errors: {r.hard_errors[:3]}" if not has_error else "")


def test_12_empty_band_section():
    """Band heading with no body content — soft warning."""
    doc = load_template()
    body = doc.element.body
    children = list(body)

    in_korean = False
    to_remove = []
    for child in children:
        if child.tag == qn("w:p"):
            p = Paragraph(child, doc)
            if p.style and p.style.style_id == "EL_BandHeading" and "Korean" in p.text:
                in_korean = True
                continue
            if in_korean:
                if p.style and p.style.style_id in ("EL_BandHeading", "EL_LegendHeading"):
                    break
                to_remove.append(child)

    for elem in to_remove:
        body.remove(elem)

    path = save_fixture(doc, "12_empty_band")
    r = parse_docx(str(path))
    has_warning = any("Korean" in w and "no body content" in w for w in r.soft_warnings)
    report("12: Empty band section warns", has_warning,
           f"Warnings: {r.soft_warnings}" if not has_warning else "")


# ============================================================
# Tests 13-15: Edge cases
# ============================================================

def test_13_cjk_adjacent_tokens():
    """CJK text with token directly adjacent — token still isolates."""
    doc = load_template()
    path = save_fixture(doc, "13_cjk_adjacent")
    r = parse_docx(str(path))

    chinese = [s for s in r.sections if s.language == "Chinese"]
    if not chinese:
        report("13: CJK adjacent tokens isolate", False, "No Chinese section found")
        return

    found = False
    for para in chinese[0].paragraphs:
        token_runs = [run for run in para.runs if run.token and run.text == "[DEVICE_MODEL]"]
        if token_runs:
            for idx, run in enumerate(para.runs):
                if run.token and run.text == "[DEVICE_MODEL]":
                    if idx > 0 and "[DEVICE_MODEL]" not in para.runs[idx - 1].text:
                        found = True
                    break

    report("13: CJK adjacent tokens isolate", found,
           "" if found else "Token not properly isolated from adjacent CJK text")


def test_14_back_to_back_tokens():
    """Two tokens back-to-back — two separate token runs."""
    doc = load_template()
    para = find_para_by_text(doc, "Hello")
    if para:
        for run_elem in para._element.findall(f".//{qn('w:t')}"):
            if "[DEVICE_USER_NAME]" in run_elem.text:
                run_elem.text = "[DEVICE_USER_NAME][DEVICE_MODEL]"
                break

    path = save_fixture(doc, "14_back_to_back_tokens")
    r = parse_docx(str(path))
    eng = [s for s in r.sections if s.language == "English"]
    if not eng:
        report("14: Back-to-back tokens split", False, "No English section")
        return

    greeting = eng[0].paragraphs[0]
    token_runs = [run for run in greeting.runs if run.token]
    two_tokens = (
        len(token_runs) >= 2
        and any(r.text == "[DEVICE_USER_NAME]" for r in token_runs)
        and any(r.text == "[DEVICE_MODEL]" for r in token_runs)
    )
    report("14: Back-to-back tokens split into separate runs", two_tokens,
           f"Token runs: {[(r.text, r.token) for r in greeting.runs]}" if not two_tokens else "")


def test_15_nested_emphasis():
    """Bold+italic vs italic-only — tracked per-run."""
    doc = load_template()
    path = save_fixture(doc, "15_nested_emphasis")
    r = parse_docx(str(path))

    eng = [s for s in r.sections if s.language == "English"]
    if not eng:
        report("15: Nested emphasis tracked", False, "No English section")
        return

    note_para = None
    for para in eng[0].paragraphs:
        if any("Note:" in run.text for run in para.runs):
            note_para = para
            break

    if note_para is None:
        report("15: Nested emphasis tracked", False, "Could not find Note: paragraph")
        return

    note_run = None
    rest_run = None
    for run in note_para.runs:
        if "Note:" in run.text:
            note_run = run
        elif "supervisor" in run.text:
            rest_run = run

    ok = (
        note_run is not None and note_run.bold and note_run.italic
        and rest_run is not None and not rest_run.bold and rest_run.italic
    )
    report("15: Nested emphasis (bold+italic vs italic) tracked", ok,
           f"Note: bold={note_run.bold if note_run else '?'} italic={note_run.italic if note_run else '?'}, "
           f"Rest: bold={rest_run.bold if rest_run else '?'} italic={rest_run.italic if rest_run else '?'}"
           if not ok else "")


# ============================================================
# Tests 16-17: Preamble & Subject
# ============================================================

def test_16_preamble_skipped():
    """Instructions preamble before LEGEND is ignored, not rejected."""
    per_lang_path = Path("/mnt/user-data/outputs/band4-english.docx")
    if not per_lang_path.exists():
        report("16: Preamble skipped cleanly", False, "Per-language template not found")
        return

    doc = Document(str(per_lang_path))
    path = save_fixture(doc, "16_preamble")
    r = parse_docx(str(path))

    ok = len(r.hard_errors) == 0 and len(r.sections) == 1
    report("16: Preamble skipped cleanly (no errors)", ok,
           f"Errors: {r.hard_errors[:3]}" if not ok else "")


def test_17_subject_parsed():
    """EL_SubjectHeading + EL_Subject parsed into section.subject."""
    per_lang_path = Path("/mnt/user-data/outputs/band4-english.docx")
    if not per_lang_path.exists():
        report("17: Subject line parsed", False, "Per-language template not found")
        return

    doc = Document(str(per_lang_path))
    path = save_fixture(doc, "17_subject")
    r = parse_docx(str(path))

    if not r.sections:
        report("17: Subject line parsed", False, "No sections found")
        return

    section = r.sections[0]
    ok = section.subject is not None and len(section.subject) > 0
    subj_in_paras = any(
        any(run.text == section.subject for run in p.runs)
        for p in section.paragraphs
    ) if section.subject else False

    ok = ok and not subj_in_paras
    report("17: Subject line parsed (not in body paragraphs)", ok,
           f"Subject: {section.subject}" if ok else f"Subject: {section.subject}, in_paras: {subj_in_paras}")


# ============================================================
# Test 18: Golden path
# ============================================================

def test_18_golden_path():
    """Full clean parse of per-language English doc — no errors, all content present."""
    per_lang_path = Path("/mnt/user-data/outputs/band4-english.docx")
    if not per_lang_path.exists():
        report("18: Golden path parse", False, "Per-language template not found")
        return

    doc = Document(str(per_lang_path))
    path = save_fixture(doc, "18_golden_path")
    r = parse_docx(str(path))

    checks = {
        "no_hard_errors": len(r.hard_errors) == 0,
        "one_section": len(r.sections) == 1,
        "band_4": r.sections[0].band == 4 if r.sections else False,
        "english": r.sections[0].language == "English" if r.sections else False,
        "has_subject": r.sections[0].subject is not None if r.sections else False,
        "has_paragraphs": len(r.sections[0].paragraphs) > 0 if r.sections else False,
        "3_dynamic_tokens": len(r.legend.dynamic_tokens) == 3,
        "3_link_tokens": len(r.legend.link_tokens) == 3,
        "tokens_isolated": any(
            any(run.token for run in p.runs)
            for p in r.sections[0].paragraphs
        ) if r.sections else False,
    }

    all_ok = all(checks.values())
    failures = [k for k, v in checks.items() if not v]
    report("18: Golden path — clean parse, all content present", all_ok,
           f"Failed: {failures}" if not all_ok else "")


# ============================================================
# Runner
# ============================================================

def main():
    print("\n" + "=" * 60)
    print("SNIFFer Parser — Smoke Tests")
    print("=" * 60)

    print("\n--- Style violations (should hard-fail) ---")
    test_01_body_to_normal()
    test_02_bullet_to_list_paragraph()
    test_03_delete_el_rule()

    print("\n--- Token violations (should hard-fail) ---")
    test_04_token_typo()
    test_05_override_new_token()

    print("\n--- Track Changes / Comments (should hard-fail) ---")
    test_06_track_changes()
    test_07_comments()

    print("\n--- Soft warnings ---")
    test_08_unused_token()
    test_09_noop_override()

    print("\n--- Structure ---")
    test_10_missing_global_legend()
    test_11_band_before_legend()
    test_12_empty_band_section()

    print("\n--- Edge cases ---")
    test_13_cjk_adjacent_tokens()
    test_14_back_to_back_tokens()
    test_15_nested_emphasis()

    print("\n--- Preamble & Subject ---")
    test_16_preamble_skipped()
    test_17_subject_parsed()

    print("\n--- Golden path ---")
    test_18_golden_path()

    print("\n" + "=" * 60)
    print(f"Results: {passed} passed, {failed} failed out of {passed + failed}")
    print("=" * 60)

    if FIXTURE_DIR.exists():
        fixtures = sorted(FIXTURE_DIR.glob("*.docx"))
        print(f"\nFixtures preserved at: {FIXTURE_DIR}/")
        print(f"  {len(fixtures)} fixture files")

    print()
    sys.exit(0 if failed == 0 else 1)


if __name__ == "__main__":
    main()
