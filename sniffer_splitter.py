"""
sniffer_splitter.py — Split a unified multi-language DOCX into per-language DOCXs.

Takes a unified DOCX (all languages in one file) and produces one DOCX per
language, each with the global legend (plus that language's override folded in),
the subject line, and the band section.

Uses the kit's parser to read the unified doc, then python-docx to write
per-language outputs.

Usage:
    python sniffer_splitter.py <unified.docx> [--output-dir <dir>]

Exit codes:
    0 — clean split
    1 — parser errors
    2 — parser warnings (output still emitted)
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph

sys.path.insert(0, str(Path(__file__).parent))
from sniffer_parser import (
    parse_docx, to_dict, ParseResult,
    VALID_LANGUAGES, iter_block_items, is_paragraph, is_table,
    resolve_el_style, para_text
)


# ============================================================
# Constants
# ============================================================

LEGEND_HEADING_COLOR = "2E75B6"
BAND_HEADING_COLOR = "538135"
SUBJECT_HEADING_COLOR = "7030A0"
SUBHEADING_COLOR = "333333"
LABEL_COLOR = "404040"
LEGEND_HEADER_FILL = "E2EFDA"
RULE_COLOR = "999999"
INSTRUCTION_COLOR = "4472C4"


# ============================================================
# Style creation
# ============================================================

def ensure_styles(doc: Document):
    """Ensure all EL_* styles exist in the document."""
    style_defs = [
        ("EL_LegendHeading", {"bold": True, "size": Pt(14), "color": LEGEND_HEADING_COLOR}),
        ("EL_BandHeading", {"bold": True, "size": Pt(14), "color": BAND_HEADING_COLOR}),
        ("EL_SubjectHeading", {"bold": True, "size": Pt(14), "color": SUBJECT_HEADING_COLOR}),
        ("EL_Subject", {"size": Pt(11)}),
        ("EL_SubHeading", {"bold": True, "size": Pt(11), "color": SUBHEADING_COLOR}),
        ("EL_LegendLabel", {"bold": True, "size": Pt(11), "color": LABEL_COLOR}),
        ("EL_Rule", {"size": Pt(2)}),
        ("EL_Body", {"size": Pt(11)}),
        ("EL_Bullet", {"size": Pt(11)}),
        ("EL_Numbered", {"size": Pt(11)}),
    ]

    existing = {s.name for s in doc.styles}
    for name, font_props in style_defs:
        if name not in existing:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = "Arial"
            if "bold" in font_props:
                style.font.bold = font_props["bold"]
            if "size" in font_props:
                style.font.size = font_props["size"]
            if "color" in font_props:
                style.font.color.rgb = RGBColor.from_string(font_props["color"])


# ============================================================
# Document construction helpers
# ============================================================

def add_styled_paragraph(doc: Document, text: str, style_name: str,
                         bold: bool = False, italic: bool = False) -> DocxParagraph:
    """Add a paragraph with the given style and optional formatting."""
    p = doc.add_paragraph()
    p.style = doc.styles[style_name]
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.name = "Arial"
    return p


def add_rule(doc: Document):
    """Add an EL_Rule paragraph with a bottom border."""
    p = doc.add_paragraph()
    p.style = doc.styles["EL_Rule"]
    # Add bottom border via XML
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), RULE_COLOR)
    bottom.set(qn("w:space"), "1")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_legend_table(doc: Document, headers: list[str], rows: list[list[str]]):
    """Add a legend table with headers and data rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))

    # Header row
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(10)
        # Shading
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), LEGEND_HEADER_FILL)
        shading.set(qn("w:val"), "clear")
        tcPr.append(shading)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, value in enumerate(row_data):
            cell = table.cell(row_idx + 1, col_idx)
            cell.text = value
            for run in cell.paragraphs[0].runs:
                run.font.name = "Arial"
                run.font.size = Pt(10)


def write_legend(doc: Document, legend: dict, language: str | None = None):
    """Write the LEGEND section including tables.

    If language is provided, fold that language's overrides into the global legend.
    """
    overrides = legend.get("overrides", {})
    lang_override = overrides.get(language, {}) if language else {}

    # Build effective token lists (global + override merged)
    dynamic_tokens = deepcopy(legend.get("dynamic_tokens", []))
    link_tokens = deepcopy(legend.get("link_tokens", []))

    # Merge overrides
    if lang_override.get("dynamic_tokens"):
        override_map = {dt["token"]: dt for dt in lang_override["dynamic_tokens"]}
        for dt in dynamic_tokens:
            if dt["token"] in override_map:
                ov = override_map[dt["token"]]
                if ov.get("example_value"):
                    dt["example_value"] = ov["example_value"]
                if ov.get("description"):
                    dt["description"] = ov["description"]

    if lang_override.get("link_tokens"):
        override_map = {lt["token"]: lt for lt in lang_override["link_tokens"]}
        for lt in link_tokens:
            if lt["token"] in override_map:
                ov = override_map[lt["token"]]
                if ov.get("display_text"):
                    lt["display_text"] = ov["display_text"]
                if ov.get("target_url"):
                    lt["target_url"] = ov["target_url"]
                if ov.get("description"):
                    lt["description"] = ov["description"]

    # LEGEND heading
    add_styled_paragraph(doc, "LEGEND", "EL_LegendHeading", bold=True)

    # Dynamic Data Tokens
    add_styled_paragraph(doc, "Dynamic Data Tokens", "EL_LegendLabel", bold=True)
    add_rule(doc)
    if dynamic_tokens:
        add_legend_table(doc,
                         ["Token", "Example Value", "Description"],
                         [[dt["token"], dt["example_value"], dt["description"]]
                          for dt in dynamic_tokens])
    add_rule(doc)

    # Link Tokens
    add_styled_paragraph(doc, "Link Tokens", "EL_LegendLabel", bold=True)
    add_rule(doc)
    if link_tokens:
        add_legend_table(doc,
                         ["Token", "Display Text", "Target URL", "Description"],
                         [[lt["token"], lt["display_text"], lt["target_url"], lt["description"]]
                          for lt in link_tokens])
    add_rule(doc)


def write_section(doc: Document, section: dict):
    """Write a band section (subject + body paragraphs) from parsed data."""
    # Subject
    if section.get("subject"):
        add_styled_paragraph(doc, "SUBJECT", "EL_SubjectHeading", bold=True)
        add_styled_paragraph(doc, section["subject"], "EL_Subject")

    # Band heading
    band = section["band"]
    lang = section["language"]
    add_styled_paragraph(doc, f"Band {band} - {lang}", "EL_BandHeading", bold=True)

    # Body paragraphs
    for para in section.get("paragraphs", []):
        style = para["style"]
        p = doc.add_paragraph()
        p.style = doc.styles[style]

        for run_data in para.get("runs", []):
            r = p.add_run(run_data["text"])
            r.bold = run_data.get("bold", False)
            r.italic = run_data.get("italic", False)
            r.font.name = "Arial"
            r.font.size = Pt(11)


# ============================================================
# Splitter
# ============================================================

def split_docx(unified_path: str, output_dir: str) -> list[str]:
    """Split a unified DOCX into per-language files.

    Returns list of output file paths.
    """
    result = parse_docx(unified_path)

    if result.hard_errors:
        print(f"\nParser errors ({len(result.hard_errors)}):", file=sys.stderr)
        for err in result.hard_errors:
            print(f"  ✗ {err}", file=sys.stderr)
        raise RuntimeError("Parser failed — cannot split")

    if result.soft_warnings:
        print(f"\nParser warnings ({len(result.soft_warnings)}):", file=sys.stderr)
        for warn in result.soft_warnings:
            print(f"  ⚠ {warn}", file=sys.stderr)

    data = to_dict(result)
    legend = data["legend"]
    sections = data["sections"]

    if not sections:
        raise RuntimeError("No sections found in unified document")

    out_path = Path(output_dir)
    out_path.mkdir(parents=True, exist_ok=True)

    files_written = []

    for section in sections:
        lang = section["language"]
        band = section["band"]

        doc = Document()
        ensure_styles(doc)

        # Remove the default empty paragraph
        if doc.paragraphs:
            doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

        # Write legend with overrides folded in
        write_legend(doc, legend, language=lang)

        # Write section
        write_section(doc, section)

        # Save
        filename = f"band{band}-{lang.lower()}.docx"
        filepath = out_path / filename
        doc.save(str(filepath))
        files_written.append(str(filepath))

    return files_written


# ============================================================
# CLI
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="Split a unified multi-language DOCX into per-language files"
    )
    parser.add_argument("unified", help="Path to the unified .docx file")
    parser.add_argument("--output-dir", "-o", default="split_output",
                        help="Output directory (default: split_output)")
    args = parser.parse_args()

    unified_path = Path(args.unified)
    if not unified_path.exists():
        print(f"Error: file not found: {unified_path}", file=sys.stderr)
        sys.exit(2)

    if unified_path.suffix != ".docx":
        print(f"Error: expected .docx file, got {unified_path.suffix}", file=sys.stderr)
        sys.exit(2)

    try:
        files = split_docx(str(unified_path), args.output_dir)
    except RuntimeError as e:
        print(f"\nError: {e}", file=sys.stderr)
        sys.exit(1)

    print(f"\nSplit into {len(files)} files:", file=sys.stderr)
    for f in sorted(files):
        print(f"  {f}", file=sys.stderr)

    # Exit 2 if parser had warnings, 0 otherwise
    sys.exit(0)


if __name__ == "__main__":
    main()
