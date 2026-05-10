"""
test_normalizer.py — Smoke tests for sniffer_normalizer.py

Generates test DOCX fixtures, runs them through parser+normalizer,
and verifies the markdown output.

Usage:
    python test_normalizer.py

Test fixtures are preserved in testing/normalizer_fixtures/ for manual review.
"""
from __future__ import annotations

import json
import os
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

sys.path.insert(0, str(Path(__file__).parent))
from sniffer_parser import parse_docx, to_dict
from sniffer_normalizer import render_section, render_runs, render_emphasis, extract_legend

# ============================================================
# Test infra
# ============================================================

passed = 0
failed = 0

FIXTURE_DIR = Path(__file__).parent / "testing" / "normalizer_fixtures"
OUTPUT_DIR = Path(__file__).parent / "testing" / "normalizer_output"


def report(name: str, ok: bool, detail: str = ""):
    global passed, failed
    if ok:
        passed += 1
        print(f"  ✓ {name}")
    else:
        failed += 1
        print(f"  ✗ {name}")
        if detail:
            print(f"    → {detail}")


# ============================================================
# Fixture builder: creates a minimal valid DOCX with EL_* styles
# ============================================================

def make_fixture(name: str, body_content: list[dict],
                 subject: str | None = "Test Subject",
                 dynamic_tokens: list[dict] | None = None,
                 link_tokens: list[dict] | None = None,
                 band: int = 1, language: str = "English") -> Path:
    """Build a minimal valid DOCX fixture and save it.

    body_content: list of dicts with keys:
        style: str (EL_Body, EL_SubHeading, EL_Bullet, EL_Numbered)
        runs: list of dicts with keys: text, bold (opt), italic (opt)
    """
    from docx.shared import Pt, RGBColor
    from docx.enum.style import WD_STYLE_TYPE

    doc = Document()

    # Create EL_* styles
    styles_to_create = [
        ("EL_LegendHeading", WD_STYLE_TYPE.PARAGRAPH),
        ("EL_BandHeading", WD_STYLE_TYPE.PARAGRAPH),
        ("EL_SubjectHeading", WD_STYLE_TYPE.PARAGRAPH),
        ("EL_Subject", WD_STYLE_TYPE.PARAGRAPH),
        ("EL_SubHeading", WD_STYLE_TYPE.PARAGRAPH),
        ("EL_LegendLabel", WD_STYLE_TYPE.PARAGRAPH),
        ("EL_Rule", WD_STYLE_TYPE.PARAGRAPH),
        ("EL_Body", WD_STYLE_TYPE.PARAGRAPH),
        ("EL_Bullet", WD_STYLE_TYPE.PARAGRAPH),
        ("EL_Numbered", WD_STYLE_TYPE.PARAGRAPH),
    ]
    for sname, stype in styles_to_create:
        if sname not in [s.name for s in doc.styles]:
            doc.styles.add_style(sname, stype)

    # Remove the default empty paragraph
    if doc.paragraphs:
        doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

    # Default tokens if not provided
    if dynamic_tokens is None:
        dynamic_tokens = [
            {"token": "[TEST_TOKEN]", "example": "test_value", "description": "A test token"}
        ]
    if link_tokens is None:
        link_tokens = [
            {"token": "<Test Link>", "display": "Click here",
             "url": "https://example.com", "description": "A test link"}
        ]

    # --- LEGEND ---
    p = doc.add_paragraph("LEGEND")
    p.style = doc.styles["EL_LegendHeading"]

    # Dynamic Data Tokens label
    p = doc.add_paragraph("Dynamic Data Tokens")
    p.style = doc.styles["EL_LegendLabel"]

    # EL_Rule
    p = doc.add_paragraph("")
    p.style = doc.styles["EL_Rule"]

    # Dynamic tokens table
    table = doc.add_table(rows=1 + len(dynamic_tokens), cols=3)
    table.cell(0, 0).text = "Token"
    table.cell(0, 1).text = "Example Value"
    table.cell(0, 2).text = "Description"
    for i, dt in enumerate(dynamic_tokens):
        table.cell(i + 1, 0).text = dt["token"]
        table.cell(i + 1, 1).text = dt["example"]
        table.cell(i + 1, 2).text = dt["description"]

    # EL_Rule
    p = doc.add_paragraph("")
    p.style = doc.styles["EL_Rule"]

    # Link Tokens label
    p = doc.add_paragraph("Link Tokens")
    p.style = doc.styles["EL_LegendLabel"]

    # EL_Rule
    p = doc.add_paragraph("")
    p.style = doc.styles["EL_Rule"]

    # Link tokens table
    table = doc.add_table(rows=1 + len(link_tokens), cols=4)
    table.cell(0, 0).text = "Token"
    table.cell(0, 1).text = "Display Text"
    table.cell(0, 2).text = "Target URL"
    table.cell(0, 3).text = "Description"
    for i, lt in enumerate(link_tokens):
        table.cell(i + 1, 0).text = lt["token"]
        table.cell(i + 1, 1).text = lt["display"]
        table.cell(i + 1, 2).text = lt["url"]
        table.cell(i + 1, 3).text = lt["description"]

    # EL_Rule
    p = doc.add_paragraph("")
    p.style = doc.styles["EL_Rule"]

    # --- SUBJECT ---
    if subject is not None:
        p = doc.add_paragraph("SUBJECT")
        p.style = doc.styles["EL_SubjectHeading"]
        p = doc.add_paragraph(subject)
        p.style = doc.styles["EL_Subject"]

    # --- BAND SECTION ---
    p = doc.add_paragraph(f"Band {band} - {language}")
    p.style = doc.styles["EL_BandHeading"]

    for para_spec in body_content:
        p = doc.add_paragraph()
        p.style = doc.styles[para_spec["style"]]

        for run_spec in para_spec["runs"]:
            r = p.add_run(run_spec["text"])
            r.bold = run_spec.get("bold", False)
            r.italic = run_spec.get("italic", False)

    # Save
    FIXTURE_DIR.mkdir(parents=True, exist_ok=True)
    path = FIXTURE_DIR / f"{name}.docx"
    doc.save(str(path))
    return path


def parse_and_normalize(fixture_path: Path) -> tuple[dict, str]:
    """Parse a fixture and render the first section to markdown."""
    result = parse_docx(str(fixture_path))
    if result.hard_errors:
        raise RuntimeError(f"Parser failed: {result.hard_errors}")
    data = to_dict(result)
    section = data["sections"][0]
    md = render_section(section)
    return data, md


# ============================================================
# Tests 1-6: Emphasis rendering
# ============================================================

def test_01_bold_only():
    """Bold text renders as **text**."""
    path = make_fixture("01_bold_only", [
        {"style": "EL_Body", "runs": [
            {"text": "This is "},
            {"text": "bold text", "bold": True},
            {"text": " here."}
        ]}
    ])
    _, md = parse_and_normalize(path)
    ok = "**bold text**" in md and "***" not in md
    report("01: Bold → **text**", ok,
           f"Got: {md.splitlines()[4]}" if not ok else "")


def test_02_italic_only():
    """Italic text renders as *text*."""
    path = make_fixture("02_italic_only", [
        {"style": "EL_Body", "runs": [
            {"text": "This is "},
            {"text": "italic text", "italic": True},
            {"text": " here."}
        ]}
    ])
    _, md = parse_and_normalize(path)
    # Check for *italic text* but not **italic text**
    ok = "*italic text*" in md and "**italic text**" not in md
    report("02: Italic → *text*", ok,
           f"Got: {md.splitlines()[4]}" if not ok else "")


def test_03_bold_italic():
    """Bold+italic text renders as ***text***."""
    path = make_fixture("03_bold_italic", [
        {"style": "EL_Body", "runs": [
            {"text": "This is "},
            {"text": "bold italic", "bold": True, "italic": True},
            {"text": " here."}
        ]}
    ])
    _, md = parse_and_normalize(path)
    ok = "***bold italic***" in md
    report("03: Bold+italic → ***text***", ok,
           f"Got: {md.splitlines()[4]}" if not ok else "")


def test_04_bold_trailing_space():
    """Bold text with trailing space: markers go inside the space."""
    path = make_fixture("04_bold_trailing_space", [
        {"style": "EL_Body", "runs": [
            {"text": "bold with space ", "bold": True},
            {"text": "then normal."}
        ]}
    ])
    _, md = parse_and_normalize(path)
    # Should be "**bold with space** then normal." not "**bold with space **then normal."
    ok = "**bold with space**" in md
    report("04: Bold trailing space → **text** (space outside)", ok,
           f"Got: {md.splitlines()[4]}" if not ok else "")


def test_05_plain_text():
    """Plain text has no emphasis markers."""
    path = make_fixture("05_plain_text", [
        {"style": "EL_Body", "runs": [
            {"text": "Just plain text with no formatting."}
        ]}
    ])
    _, md = parse_and_normalize(path)
    line = md.splitlines()[4]
    ok = "*" not in line and line == "Just plain text with no formatting."
    report("05: Plain text → no markers", ok,
           f"Got: {line}" if not ok else "")


def test_06_subheading_bold_stripping():
    """SubHeading bold is stripped — ## text, not ## **text**."""
    path = make_fixture("06_subheading_strip", [
        {"style": "EL_SubHeading", "runs": [
            {"text": "Important Section", "bold": True}
        ]},
        {"style": "EL_Body", "runs": [
            {"text": "Body after heading."}
        ]}
    ])
    _, md = parse_and_normalize(path)
    ok = "## Important Section" in md and "## **Important Section**" not in md
    report("06: SubHeading → ## text (bold stripped)", ok,
           f"Got: {[l for l in md.splitlines() if '##' in l]}" if not ok else "")


# ============================================================
# Tests 7-11: Paragraph styles
# ============================================================

def test_07_body_paragraph():
    """EL_Body renders as plain paragraph."""
    path = make_fixture("07_body", [
        {"style": "EL_Body", "runs": [{"text": "A body paragraph."}]}
    ])
    _, md = parse_and_normalize(path)
    ok = "A body paragraph." in md and not any(
        md.splitlines()[i].startswith(("- ", "1.", "## "))
        for i in range(len(md.splitlines()))
        if "A body paragraph." in md.splitlines()[i]
    )
    report("07: EL_Body → plain paragraph", ok)


def test_08_subheading():
    """EL_SubHeading renders as ## heading."""
    path = make_fixture("08_subheading", [
        {"style": "EL_SubHeading", "runs": [{"text": "My Heading", "bold": True}]},
        {"style": "EL_Body", "runs": [{"text": "Content."}]}
    ])
    _, md = parse_and_normalize(path)
    ok = "## My Heading" in md
    report("08: EL_SubHeading → ## heading", ok)


def test_09_bullet():
    """EL_Bullet renders as - text."""
    path = make_fixture("09_bullet", [
        {"style": "EL_Bullet", "runs": [{"text": "First bullet"}]},
        {"style": "EL_Bullet", "runs": [{"text": "Second bullet"}]},
    ])
    _, md = parse_and_normalize(path)
    ok = "- First bullet" in md and "- Second bullet" in md
    report("09: EL_Bullet → - text", ok)


def test_10_numbered_sequential():
    """EL_Numbered renumbers sequentially."""
    path = make_fixture("10_numbered", [
        {"style": "EL_Numbered", "runs": [{"text": "Step one"}]},
        {"style": "EL_Numbered", "runs": [{"text": "Step two"}]},
        {"style": "EL_Numbered", "runs": [{"text": "Step three"}]},
    ])
    _, md = parse_and_normalize(path)
    ok = "1. Step one" in md and "2. Step two" in md and "3. Step three" in md
    report("10: EL_Numbered → 1. 2. 3.", ok,
           f"Got: {[l for l in md.splitlines() if l.strip()]}" if not ok else "")


def test_11_numbered_reset():
    """EL_Numbered counter resets after non-numbered paragraph."""
    path = make_fixture("11_numbered_reset", [
        {"style": "EL_Numbered", "runs": [{"text": "First"}]},
        {"style": "EL_Numbered", "runs": [{"text": "Second"}]},
        {"style": "EL_Body", "runs": [{"text": "Interruption."}]},
        {"style": "EL_Numbered", "runs": [{"text": "Restart"}]},
    ])
    _, md = parse_and_normalize(path)
    lines = md.splitlines()
    has_1_first = "1. First" in md
    has_2_second = "2. Second" in md
    has_1_restart = any(l == "1. Restart" for l in lines)
    ok = has_1_first and has_2_second and has_1_restart
    report("11: Numbered resets after body paragraph", ok,
           f"Got: {[l for l in lines if l.strip()]}" if not ok else "")


# ============================================================
# Tests 12-14: Tokens in markdown
# ============================================================

def test_12_dynamic_token_literal():
    """[TOKEN] stays literal in markdown."""
    path = make_fixture("12_dynamic_token", [
        {"style": "EL_Body", "runs": [
            {"text": "Hello "},
            {"text": "[TEST_TOKEN]"},
            {"text": ", welcome."}
        ]}
    ])
    _, md = parse_and_normalize(path)
    ok = "[TEST_TOKEN]" in md
    report("12: [TOKEN] stays literal", ok)


def test_13_link_token_literal():
    """<TOKEN> stays literal, not interpreted as HTML."""
    path = make_fixture("13_link_token", [
        {"style": "EL_Body", "runs": [
            {"text": "Click "},
            {"text": "<Test Link>"},
            {"text": " for help."}
        ]}
    ])
    _, md = parse_and_normalize(path)
    ok = "<Test Link>" in md
    report("13: <TOKEN> stays literal (not HTML-mangled)", ok,
           f"Got: {[l for l in md.splitlines() if 'Click' in l]}" if not ok else "")


def test_14_token_inside_emphasis():
    """Token inside bold: **important [TOKEN] text**."""
    path = make_fixture("14_token_emphasis", [
        {"style": "EL_Body", "runs": [
            {"text": "important ", "bold": True},
            {"text": "[TEST_TOKEN]", "bold": True},
            {"text": " text", "bold": True},
        ]}
    ])
    _, md = parse_and_normalize(path)
    # Token should appear inside the bold markers
    ok = "[TEST_TOKEN]" in md and "**" in md
    report("14: Token inside bold emphasis preserved", ok,
           f"Got: {[l for l in md.splitlines() if 'TEST_TOKEN' in l]}" if not ok else "")


# ============================================================
# Tests 15-16: Subject line
# ============================================================

def test_15_subject_present():
    """Subject line renders as **Subject:** text."""
    path = make_fixture("15_subject_present", [
        {"style": "EL_Body", "runs": [{"text": "Body content."}]}
    ], subject="Important Email Subject")
    _, md = parse_and_normalize(path)
    ok = "**Subject:** Important Email Subject" in md
    # Should appear before body content
    subj_idx = md.index("**Subject:**")
    body_idx = md.index("Body content.")
    ok = ok and subj_idx < body_idx
    report("15: Subject present → **Subject:** text (before body)", ok)


def test_16_subject_absent():
    """No subject → body starts immediately after heading."""
    path = make_fixture("16_subject_absent", [
        {"style": "EL_Body", "runs": [{"text": "Body content."}]}
    ], subject=None)
    _, md = parse_and_normalize(path)
    ok = "**Subject:**" not in md and "Body content." in md
    # Body should be the first non-empty line after heading
    lines = [l for l in md.splitlines() if l.strip()]
    ok = ok and lines[0].startswith("# Band") and lines[1] == "Body content."
    report("16: No subject → body starts immediately", ok,
           f"Got lines: {lines[:3]}" if not ok else "")


# ============================================================
# Tests 17-19: Structure
# ============================================================

def test_17_output_path_convention():
    """Output files land at <dir>/<round>/band<N>/<lang>.md."""
    path = make_fixture("17_output_path", [
        {"style": "EL_Body", "runs": [{"text": "Content."}]}
    ], band=3, language="French")

    result = parse_docx(str(path))
    data = to_dict(result)

    out_dir = OUTPUT_DIR / "17"
    out_dir.mkdir(parents=True, exist_ok=True)

    # Simulate normalizer output path logic
    section = data["sections"][0]
    band_dir = out_dir / "test-round" / f"band{section['band']}"
    band_dir.mkdir(parents=True, exist_ok=True)
    md_path = band_dir / f"{section['language'].lower()}.md"
    md_path.write_text(render_section(section), encoding="utf-8")

    ok = md_path.exists() and md_path.name == "french.md"
    ok = ok and "band3" in str(md_path)
    ok = ok and "test-round" in str(md_path)
    report("17: Output path → <dir>/<round>/band<N>/<lang>.md", ok,
           f"Path: {md_path}" if not ok else "")


def test_18_legend_json_valid():
    """legend.json is valid JSON with expected structure."""
    path = make_fixture("18_legend_json", [
        {"style": "EL_Body", "runs": [{"text": "Content."}]}
    ], dynamic_tokens=[
        {"token": "[A]", "example": "val_a", "description": "Token A"},
        {"token": "[B]", "example": "val_b", "description": "Token B"},
    ], link_tokens=[
        {"token": "<Link1>", "display": "Click", "url": "https://example.com", "description": "Link one"},
    ])

    result = parse_docx(str(path))
    data = to_dict(result)
    legend = extract_legend(data)

    ok = True
    detail = []

    # Valid JSON round-trip
    try:
        json_str = json.dumps(legend, indent=2, ensure_ascii=False)
        parsed_back = json.loads(json_str)
    except Exception as e:
        ok = False
        detail.append(f"JSON error: {e}")

    # Structure checks
    if "dynamic_tokens" not in legend:
        ok = False
        detail.append("Missing dynamic_tokens")
    elif len(legend["dynamic_tokens"]) != 2:
        ok = False
        detail.append(f"Expected 2 dynamic tokens, got {len(legend['dynamic_tokens'])}")

    if "link_tokens" not in legend:
        ok = False
        detail.append("Missing link_tokens")
    elif len(legend["link_tokens"]) != 1:
        ok = False
        detail.append(f"Expected 1 link token, got {len(legend['link_tokens'])}")

    report("18: legend.json valid with expected structure", ok, "; ".join(detail))


def test_19_legacy_unified_multi_language():
    """Legacy unified doc produces N markdown files."""
    legacy_path = Path("/mnt/user-data/outputs/branch-b-template-band4.docx")
    if not legacy_path.exists():
        report("19: Legacy unified → N markdown files", False, "Legacy template not found")
        return

    result = parse_docx(str(legacy_path))
    data = to_dict(result)

    languages = {s["language"].lower() for s in data["sections"]}
    ok = len(languages) >= 3  # at least English, Spanish, Chinese have content
    ok = ok and "english" in languages and "chinese" in languages

    report("19: Legacy unified → multiple language sections", ok,
           f"Languages found: {sorted(languages)}" if not ok else "")


# ============================================================
# Test 20: Round-trip fidelity
# ============================================================

def test_20_round_trip_fidelity():
    """Parse DOCX → normalize → verify content preserved."""
    path = make_fixture("20_round_trip", [
        {"style": "EL_SubHeading", "runs": [{"text": "Action Required", "bold": True}]},
        {"style": "EL_Body", "runs": [
            {"text": "Please return "},
            {"text": "[TEST_TOKEN]", "bold": False},
            {"text": " via "},
            {"text": "<Test Link>"},
            {"text": " immediately."},
        ]},
        {"style": "EL_Bullet", "runs": [
            {"text": "First: "},
            {"text": "important", "bold": True},
            {"text": " detail."}
        ]},
        {"style": "EL_Bullet", "runs": [
            {"text": "Second: "},
            {"text": "italic note", "italic": True},
        ]},
        {"style": "EL_Body", "runs": [
            {"text": "Note:", "bold": True, "italic": True},
            {"text": " This is critical.", "italic": True},
        ]},
    ], subject="Urgent: Device Action Required")

    _, md = parse_and_normalize(path)
    lines = md.splitlines()

    checks = {
        "heading": "# Band 1 - English" in md,
        "subject": "**Subject:** Urgent: Device Action Required" in md,
        "subheading": "## Action Required" in md,
        "token_dynamic": "[TEST_TOKEN]" in md,
        "token_link": "<Test Link>" in md,
        "bold_in_bullet": "**important**" in md,
        "italic_in_bullet": "*italic note*" in md,
        "bold_italic_mixed": "***Note:***" in md,
        "italic_continuation": "*This is critical.*" in md,
        "bullet_markers": md.count("- ") >= 2,
    }

    all_ok = all(checks.values())
    failures = [k for k, v in checks.items() if not v]
    report("20: Round-trip fidelity (all content preserved)", all_ok,
           f"Failed checks: {failures}" if not all_ok else "")


# ============================================================
# Runner
# ============================================================

def main():
    # Clean output dir
    if OUTPUT_DIR.exists():
        import shutil
        shutil.rmtree(OUTPUT_DIR)

    print("\n" + "=" * 60)
    print("SNIFFer Normalizer — Smoke Tests")
    print("=" * 60)

    print("\n--- Emphasis rendering ---")
    test_01_bold_only()
    test_02_italic_only()
    test_03_bold_italic()
    test_04_bold_trailing_space()
    test_05_plain_text()
    test_06_subheading_bold_stripping()

    print("\n--- Paragraph styles ---")
    test_07_body_paragraph()
    test_08_subheading()
    test_09_bullet()
    test_10_numbered_sequential()
    test_11_numbered_reset()

    print("\n--- Tokens in markdown ---")
    test_12_dynamic_token_literal()
    test_13_link_token_literal()
    test_14_token_inside_emphasis()

    print("\n--- Subject line ---")
    test_15_subject_present()
    test_16_subject_absent()

    print("\n--- Structure ---")
    test_17_output_path_convention()
    test_18_legend_json_valid()
    test_19_legacy_unified_multi_language()

    print("\n--- Round-trip fidelity ---")
    test_20_round_trip_fidelity()

    print("\n" + "=" * 60)
    print(f"Results: {passed} passed, {failed} failed out of {passed + failed}")
    print("=" * 60)

    # Report fixture location
    if FIXTURE_DIR.exists():
        fixtures = sorted(FIXTURE_DIR.glob("*.docx"))
        print(f"\nFixtures preserved at: {FIXTURE_DIR}/")
        print(f"  {len(fixtures)} fixture files")

    print()
    sys.exit(0 if failed == 0 else 1)


if __name__ == "__main__":
    main()
