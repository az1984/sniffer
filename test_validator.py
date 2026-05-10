"""
test_validator.py — Smoke tests for sniffer_validator.py

Generates test fixtures (DOCX and JSON), runs them through the validator,
and verifies hard errors and soft warnings.

Usage:
    python test_validator.py

Test fixtures are preserved in testing/validator_fixtures/ for manual review.
"""
from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

sys.path.insert(0, str(Path(__file__).parent))
from sniffer_parser import parse_docx, to_dict
from sniffer_validator import validate, merge_parsed

# ============================================================
# Infra
# ============================================================

passed = 0
failed = 0

FIXTURE_DIR = Path(__file__).parent / "testing" / "validator_fixtures"


def report(name: str, ok: bool, detail: str = ""):
    global passed, failed
    if ok:
        passed += 1
        print(f"  ✓ {name}")
    else:
        failed += 1
        print(f"  ✗ {name}")
        if detail:
            print(f"    → {detail}")


# ============================================================
# Fixture builder
# ============================================================

def make_fixture_docx(name: str, band: int, language: str,
                      subject: str | None = "Test Subject",
                      body_texts: list[str] | None = None,
                      dynamic_tokens: list[dict] | None = None,
                      link_tokens: list[dict] | None = None) -> Path:
    """Build a minimal valid per-language DOCX."""
    doc = Document()

    for sname in ["EL_LegendHeading", "EL_BandHeading", "EL_SubjectHeading",
                   "EL_Subject", "EL_SubHeading", "EL_LegendLabel", "EL_Rule",
                   "EL_Body", "EL_Bullet", "EL_Numbered"]:
        if sname not in [s.name for s in doc.styles]:
            doc.styles.add_style(sname, WD_STYLE_TYPE.PARAGRAPH)

    if doc.paragraphs:
        doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

    if dynamic_tokens is None:
        dynamic_tokens = [
            {"token": "[NAME]", "example": "Jane", "description": "User name"}
        ]
    if link_tokens is None:
        link_tokens = [
            {"token": "<Help>", "display": "Help Center",
             "url": "https://help.example.com", "description": "Support link"}
        ]
    if body_texts is None:
        body_texts = ["Hello [NAME], please visit <Help> for assistance."]

    # LEGEND
    p = doc.add_paragraph("LEGEND")
    p.style = doc.styles["EL_LegendHeading"]

    p = doc.add_paragraph("Dynamic Data Tokens")
    p.style = doc.styles["EL_LegendLabel"]
    p = doc.add_paragraph("")
    p.style = doc.styles["EL_Rule"]

    table = doc.add_table(rows=1 + len(dynamic_tokens), cols=3)
    table.cell(0, 0).text = "Token"
    table.cell(0, 1).text = "Example Value"
    table.cell(0, 2).text = "Description"
    for i, dt in enumerate(dynamic_tokens):
        table.cell(i + 1, 0).text = dt["token"]
        table.cell(i + 1, 1).text = dt["example"]
        table.cell(i + 1, 2).text = dt["description"]

    p = doc.add_paragraph("")
    p.style = doc.styles["EL_Rule"]

    p = doc.add_paragraph("Link Tokens")
    p.style = doc.styles["EL_LegendLabel"]
    p = doc.add_paragraph("")
    p.style = doc.styles["EL_Rule"]

    table = doc.add_table(rows=1 + len(link_tokens), cols=4)
    table.cell(0, 0).text = "Token"
    table.cell(0, 1).text = "Display Text"
    table.cell(0, 2).text = "Target URL"
    table.cell(0, 3).text = "Description"
    for i, lt in enumerate(link_tokens):
        table.cell(i + 1, 0).text = lt["token"]
        table.cell(i + 1, 1).text = lt["display"]
        table.cell(i + 1, 2).text = lt["url"]
        table.cell(i + 1, 3).text = lt["description"]

    p = doc.add_paragraph("")
    p.style = doc.styles["EL_Rule"]

    # SUBJECT
    if subject is not None:
        p = doc.add_paragraph("SUBJECT")
        p.style = doc.styles["EL_SubjectHeading"]
        p = doc.add_paragraph(subject)
        p.style = doc.styles["EL_Subject"]

    # BAND
    p = doc.add_paragraph(f"Band {band} - {language}")
    p.style = doc.styles["EL_BandHeading"]

    for text in body_texts:
        p = doc.add_paragraph(text)
        p.style = doc.styles["EL_Body"]

    FIXTURE_DIR.mkdir(parents=True, exist_ok=True)
    path = FIXTURE_DIR / f"{name}.docx"
    doc.save(str(path))
    return path


def parse_fixture(path: Path) -> dict:
    """Parse a fixture and return the dict."""
    result = parse_docx(str(path))
    if result.hard_errors:
        raise RuntimeError(f"Parser failed on {path.name}: {result.hard_errors}")
    return to_dict(result)


def make_full_language_set(name_prefix: str, **kwargs) -> list[dict]:
    """Create fixtures for all 9 languages and return parsed data."""
    languages = [
        "English", "French", "German", "Italian", "Spanish",
        "Portuguese", "Japanese", "Chinese", "Korean"
    ]
    datasets = []
    for lang in languages:
        path = make_fixture_docx(
            f"{name_prefix}_{lang.lower()}", band=1, language=lang, **kwargs
        )
        datasets.append(parse_fixture(path))
    return datasets


# ============================================================
# Tests 01-05: Hard rules
# ============================================================

def test_01_missing_languages():
    """Only 3 of 9 languages present — hard error."""
    datasets = []
    for lang in ["English", "French", "German"]:
        path = make_fixture_docx(f"01_missing_{lang.lower()}", band=1, language=lang)
        datasets.append(parse_fixture(path))

    merged = merge_parsed(datasets)
    errors, warnings = validate(merged)

    ok = any("Missing languages" in e for e in errors)
    report("01: Missing languages → hard error", ok,
           f"Errors: {errors}" if not ok else "")


def test_02_empty_example_value():
    """Dynamic token with empty example value — hard error."""
    path = make_fixture_docx("02_empty_example", band=1, language="English",
                             dynamic_tokens=[
                                 {"token": "[NAME]", "example": "", "description": "User name"},
                             ])
    data = parse_fixture(path)
    errors, warnings = validate(data)

    ok = any("empty example value" in e for e in errors)
    report("02: Empty example value → hard error", ok,
           f"Errors: {errors}" if not ok else "")


def test_03_empty_display_text():
    """Link token with empty display text — hard error."""
    path = make_fixture_docx("03_empty_display", band=1, language="English",
                             link_tokens=[
                                 {"token": "<Help>", "display": "",
                                  "url": "https://example.com", "description": "Help"},
                             ])
    data = parse_fixture(path)
    errors, warnings = validate(data)

    ok = any("empty display text" in e for e in errors)
    report("03: Empty display text → hard error", ok,
           f"Errors: {errors}" if not ok else "")


def test_04_invalid_url():
    """Link token with invalid URL — hard error."""
    path = make_fixture_docx("04_invalid_url", band=1, language="English",
                             link_tokens=[
                                 {"token": "<Help>", "display": "Help",
                                  "url": "not-a-url", "description": "Help"},
                             ])
    data = parse_fixture(path)
    errors, warnings = validate(data)

    ok = any("invalid URL" in e for e in errors)
    report("04: Invalid URL → hard error", ok,
           f"Errors: {errors}" if not ok else "")


def test_05_inconsistent_band_numbers():
    """Sections with different band numbers — hard error."""
    path_a = make_fixture_docx("05_band_a", band=1, language="English")
    path_b = make_fixture_docx("05_band_b", band=2, language="French")

    data_a = parse_fixture(path_a)
    data_b = parse_fixture(path_b)
    merged = merge_parsed([data_a, data_b])
    errors, warnings = validate(merged)

    ok = any("Inconsistent band numbers" in e for e in errors)
    report("05: Inconsistent band numbers → hard error", ok,
           f"Errors: {errors}" if not ok else "")


# ============================================================
# Tests 06-08: Soft warnings
# ============================================================

def test_06_paragraph_asymmetry():
    """One language has far fewer paragraphs — soft warning."""
    datasets = []
    for lang in ["English", "French"]:
        body = ["Line one.", "Line two.", "Line three.", "Line four.", "Line five."]
        if lang == "French":
            body = ["Ligne un."]  # only 1 paragraph vs 5
        path = make_fixture_docx(f"06_asymmetry_{lang.lower()}", band=1, language=lang,
                                 body_texts=body)
        datasets.append(parse_fixture(path))

    merged = merge_parsed(datasets)
    errors, warnings = validate(merged)

    ok = any("asymmetry" in w.lower() or "incomplete" in w.lower() for w in warnings)
    report("06: Paragraph count asymmetry → soft warning", ok,
           f"Warnings: {warnings}" if not ok else "")


def test_07_token_usage_asymmetry():
    """Token used in English but not French — soft warning."""
    tokens_d = [{"token": "[NAME]", "example": "Jane", "description": "User name"}]
    tokens_l = [{"token": "<Help>", "display": "Help", "url": "https://example.com", "description": "Help"}]

    path_en = make_fixture_docx("07_token_asym_en", band=1, language="English",
                                dynamic_tokens=tokens_d, link_tokens=tokens_l,
                                body_texts=["Hello [NAME], visit <Help>."])
    path_fr = make_fixture_docx("07_token_asym_fr", band=1, language="French",
                                dynamic_tokens=tokens_d, link_tokens=tokens_l,
                                body_texts=["Bonjour [NAME], bienvenue."])  # has [NAME] but not <Help>

    data_en = parse_fixture(path_en)
    data_fr = parse_fixture(path_fr)
    merged = merge_parsed([data_en, data_fr])
    errors, warnings = validate(merged)

    ok = any("<Help>" in w and "French" in w for w in warnings)
    report("07: Token usage asymmetry → soft warning", ok,
           f"Warnings: {warnings}" if not ok else "")


def test_08_missing_subject():
    """Section with body content but no subject — soft warning."""
    path = make_fixture_docx("08_no_subject", band=1, language="English",
                             subject=None)
    data = parse_fixture(path)
    errors, warnings = validate(data)

    ok = any("no subject" in w.lower() for w in warnings)
    report("08: Missing subject line → soft warning", ok,
           f"Warnings: {warnings}" if not ok else "")


# ============================================================
# Tests 09-11: Clean passes
# ============================================================

def test_09_single_file_clean():
    """Single well-formed file — should pass (with missing-languages error only)."""
    path = make_fixture_docx("09_single_clean", band=1, language="English")
    data = parse_fixture(path)
    errors, warnings = validate(data)

    # Only error should be missing languages
    non_lang_errors = [e for e in errors if "Missing languages" not in e]
    ok = len(non_lang_errors) == 0
    report("09: Single clean file — no non-language errors", ok,
           f"Other errors: {non_lang_errors}" if not ok else "")


def test_10_all_languages_clean():
    """All 9 languages, well-formed — should pass clean."""
    datasets = make_full_language_set("10_all_clean")
    merged = merge_parsed(datasets)
    errors, warnings = validate(merged)

    ok = len(errors) == 0 and len(warnings) == 0
    report("10: All 9 languages clean — no errors, no warnings", ok,
           f"Errors: {errors}, Warnings: {warnings}" if not ok else "")


def test_11_valid_urls():
    """Link tokens with valid URLs — no URL errors."""
    path = make_fixture_docx("11_valid_urls", band=1, language="English",
                             link_tokens=[
                                 {"token": "<Portal>", "display": "Portal",
                                  "url": "https://portal.example.com/path?q=1",
                                  "description": "Main portal"},
                                 {"token": "<Docs>", "display": "Documentation",
                                  "url": "https://docs.example.com",
                                  "description": "Documentation site"},
                             ],
                             body_texts=["Visit <Portal> or <Docs>."])
    data = parse_fixture(path)
    errors, warnings = validate(data)

    url_errors = [e for e in errors if "URL" in e]
    ok = len(url_errors) == 0
    report("11: Valid URLs — no URL errors", ok,
           f"URL errors: {url_errors}" if not ok else "")


# ============================================================
# Tests 12-13: Edge cases
# ============================================================

def test_12_empty_url():
    """Link token with completely empty URL — hard error."""
    path = make_fixture_docx("12_empty_url", band=1, language="English",
                             link_tokens=[
                                 {"token": "<Help>", "display": "Help",
                                  "url": "", "description": "Help"},
                             ])
    data = parse_fixture(path)
    errors, warnings = validate(data)

    ok = any("empty target URL" in e for e in errors)
    report("12: Empty URL → hard error", ok,
           f"Errors: {errors}" if not ok else "")


def test_13_merge_preserves_legend():
    """Merging multiple files preserves the legend from the first."""
    tokens_d = [
        {"token": "[A]", "example": "val_a", "description": "Token A"},
        {"token": "[B]", "example": "val_b", "description": "Token B"},
    ]
    tokens_l = [
        {"token": "<Help>", "display": "Help", "url": "https://example.com", "description": "Help"}
    ]

    path_en = make_fixture_docx("13_merge_en", band=1, language="English",
                                dynamic_tokens=tokens_d, link_tokens=tokens_l,
                                body_texts=["Use [A] and [B]."])
    path_fr = make_fixture_docx("13_merge_fr", band=1, language="French",
                                dynamic_tokens=tokens_d, link_tokens=tokens_l,
                                body_texts=["Utilisez [A] et [B]."])

    data_en = parse_fixture(path_en)
    data_fr = parse_fixture(path_fr)
    merged = merge_parsed([data_en, data_fr])

    ok = len(merged["legend"]["dynamic_tokens"]) == 2
    ok = ok and len(merged["sections"]) == 2
    report("13: Merge preserves legend and combines sections", ok,
           f"Tokens: {len(merged['legend']['dynamic_tokens'])}, "
           f"Sections: {len(merged['sections'])}" if not ok else "")


# ============================================================
# Runner
# ============================================================

def main():
    print("\n" + "=" * 60)
    print("SNIFFer Validator — Smoke Tests")
    print("=" * 60)

    print("\n--- Hard rules ---")
    test_01_missing_languages()
    test_02_empty_example_value()
    test_03_empty_display_text()
    test_04_invalid_url()
    test_05_inconsistent_band_numbers()

    print("\n--- Soft warnings ---")
    test_06_paragraph_asymmetry()
    test_07_token_usage_asymmetry()
    test_08_missing_subject()

    print("\n--- Clean passes ---")
    test_09_single_file_clean()
    test_10_all_languages_clean()
    test_11_valid_urls()

    print("\n--- Edge cases ---")
    test_12_empty_url()
    test_13_merge_preserves_legend()

    print("\n" + "=" * 60)
    print(f"Results: {passed} passed, {failed} failed out of {passed + failed}")
    print("=" * 60)

    if FIXTURE_DIR.exists():
        fixtures = sorted(FIXTURE_DIR.glob("*.docx"))
        print(f"\nFixtures preserved at: {FIXTURE_DIR}/")
        print(f"  {len(fixtures)} fixture files")

    print()
    sys.exit(0 if failed == 0 else 1)


if __name__ == "__main__":
    main()
