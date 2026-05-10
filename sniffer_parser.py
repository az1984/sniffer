"""
sniffer_parser.py — Parse a SNIFFer DOCX into structured runs.

Usage:
    python sniffer_parser.py <docx_path> [--output <json_path>]

Exit codes:
    0 — clean parse, output emitted
    1 — hard-rule violations (no output emitted)
    2 — soft warnings only (output emitted, warnings printed to stderr)
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml.ns import qn


# ============================================================
# Data structures
# ============================================================

@dataclass
class Run:
    text: str
    bold: bool = False
    italic: bool = False
    token: bool = False  # True if this run is a [DYNAMIC] or <LINK> token

@dataclass
class ParsedParagraph:
    style: str  # EL_Body, EL_SubHeading, EL_Bullet, EL_Numbered
    runs: list[Run] = field(default_factory=list)

@dataclass
class DynamicToken:
    token: str
    example_value: str
    description: str

@dataclass
class LinkToken:
    token: str
    display_text: str
    target_url: str
    description: str

@dataclass
class Legend:
    dynamic_tokens: list[DynamicToken] = field(default_factory=list)
    link_tokens: list[LinkToken] = field(default_factory=list)

@dataclass
class LegendOverride:
    language: str
    dynamic_tokens: list[DynamicToken] = field(default_factory=list)
    link_tokens: list[LinkToken] = field(default_factory=list)

@dataclass
class BandSection:
    band: int
    language: str
    subject: Optional[str] = None
    paragraphs: list[ParsedParagraph] = field(default_factory=list)

@dataclass
class ParseResult:
    legend: Legend
    overrides: list[LegendOverride]
    sections: list[BandSection]
    hard_errors: list[str] = field(default_factory=list)
    soft_warnings: list[str] = field(default_factory=list)


# ============================================================
# Constants
# ============================================================

VALID_LANGUAGES = {
    "English", "French", "German", "Italian", "Spanish",
    "Portuguese", "Japanese", "Chinese", "Korean"
}

EL_STYLES = {
    "EL_LegendHeading", "EL_BandHeading", "EL_SubjectHeading", "EL_Subject",
    "EL_SubHeading", "EL_LegendLabel", "EL_Rule", "EL_Body", "EL_Bullet", "EL_Numbered"
}

BAND_STYLES = {"EL_Body", "EL_SubHeading", "EL_Bullet", "EL_Numbered"}

SECTION_HEADING_STYLES = {"EL_LegendHeading", "EL_BandHeading", "EL_SubjectHeading"}

LEGEND_HEADING_RE = re.compile(r"^LEGEND(?:\s*-\s*(.+))?$")
BAND_HEADING_RE = re.compile(r"^Band\s+(\d+)\s*-\s*(.+)$")
SUBJECT_HEADING_RE = re.compile(r"^SUBJECT$")

# Token patterns — used for splitting runs, not downstream matching
TOKEN_PATTERN = re.compile(r"(\[[A-Z_]+\]|<[^>]+>)")


# ============================================================
# Helpers
# ============================================================

def get_style_name(para) -> Optional[str]:
    """Get the paragraph style name, resolving through style hierarchy."""
    style = para.style
    if style is None:
        return None
    return style.name


def get_style_id(para) -> Optional[str]:
    """Get the paragraph style ID (what's in the XML)."""
    style = para.style
    if style is None:
        return None
    return style.style_id


def is_el_style(para) -> bool:
    """Check if paragraph has an EL_* style by ID or name."""
    sid = get_style_id(para)
    sname = get_style_name(para)
    return (sid and sid.startswith("EL_")) or (sname and sname.startswith("EL_"))


def resolve_el_style(para) -> Optional[str]:
    """Return the EL_* style identifier for this paragraph, or None."""
    sid = get_style_id(para)
    if sid and sid.startswith("EL_"):
        return sid
    sname = get_style_name(para)
    if sname and sname.startswith("EL_"):
        return sname
    return None


def has_track_changes(doc) -> bool:
    """Check if document contains tracked changes."""
    body = doc.element.body
    return (
        len(body.findall(f".//{qn('w:ins')}")) > 0
        or len(body.findall(f".//{qn('w:del')}")) > 0
    )


def has_comments(doc) -> bool:
    """Check if document contains comments."""
    body = doc.element.body
    return len(body.findall(f".//{qn('w:commentRangeStart')}")) > 0


def para_text(para) -> str:
    """Get full text of a paragraph."""
    return para.text.strip()


def is_empty_para(para) -> bool:
    """Check if paragraph is effectively empty."""
    return len(para.text.strip()) == 0


def run_is_bold(run) -> bool:
    """Determine if a run is bold."""
    if run.bold is True:
        return True
    if run.bold is None and run.style and run.style.font.bold:
        return True
    return False


def run_is_italic(run) -> bool:
    """Determine if a run is italic."""
    if run.italic is True:
        return True
    if run.italic is None and run.style and run.style.font.italic:
        return True
    return False


# ============================================================
# Run processing: merge then split at tokens
# ============================================================

def extract_raw_runs(para) -> list[Run]:
    """Extract runs from a paragraph with bold/italic info."""
    raw = []
    for r in para.runs:
        text = r.text
        if not text:
            continue
        raw.append(Run(
            text=text,
            bold=run_is_bold(r),
            italic=run_is_italic(r),
            token=False
        ))
    return raw


def merge_runs(runs: list[Run]) -> list[Run]:
    """Merge adjacent runs with identical formatting."""
    if not runs:
        return []
    merged = [Run(text=runs[0].text, bold=runs[0].bold,
                  italic=runs[0].italic, token=runs[0].token)]
    for r in runs[1:]:
        prev = merged[-1]
        if prev.bold == r.bold and prev.italic == r.italic and not prev.token and not r.token:
            prev.text += r.text
        else:
            merged.append(Run(text=r.text, bold=r.bold,
                              italic=r.italic, token=r.token))
    return merged


def split_tokens(runs: list[Run]) -> list[Run]:
    """Split runs at token boundaries. Tokens get their own Run with token=True."""
    result = []
    for r in runs:
        parts = TOKEN_PATTERN.split(r.text)
        for part in parts:
            if not part:
                continue
            is_tok = bool(TOKEN_PATTERN.fullmatch(part))
            result.append(Run(
                text=part,
                bold=r.bold,
                italic=r.italic,
                token=is_tok
            ))
    return result


def process_runs(para) -> list[Run]:
    """Full pipeline: extract → merge → token-split."""
    raw = extract_raw_runs(para)
    merged = merge_runs(raw)
    split = split_tokens(merged)
    # Final merge pass (token split may have created adjacent same-format non-token runs)
    return merge_runs(split)


# ============================================================
# Table parsing
# ============================================================

def read_table_rows(table) -> list[list[str]]:
    """Read a table into a list of rows, each a list of cell text."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)
    return rows


def parse_dynamic_token_table(rows: list[list[str]], errors: list[str],
                               context: str) -> list[DynamicToken]:
    """Parse a Dynamic Data Tokens table (3 columns: token, example, description)."""
    tokens = []
    for i, row in enumerate(rows):
        if i == 0:  # header row
            if len(row) < 3:
                errors.append(f"{context}: Dynamic Data Tokens header row has {len(row)} columns, expected 3")
            continue
        if len(row) < 3:
            errors.append(f"{context}: Dynamic Data Tokens row {i} has {len(row)} columns, expected 3")
            continue
        tokens.append(DynamicToken(
            token=row[0],
            example_value=row[1],
            description=row[2]
        ))
    return tokens


def parse_link_token_table(rows: list[list[str]], errors: list[str],
                            context: str) -> list[LinkToken]:
    """Parse a Link Tokens table (4 columns: token, display, url, description)."""
    tokens = []
    for i, row in enumerate(rows):
        if i == 0:  # header row
            if len(row) < 4:
                errors.append(f"{context}: Link Tokens header row has {len(row)} columns, expected 4")
            continue
        if len(row) < 4:
            errors.append(f"{context}: Link Tokens row {i} has {len(row)} columns, expected 4")
            continue
        tokens.append(LinkToken(
            token=row[0],
            display_text=row[1],
            target_url=row[2],
            description=row[3]
        ))
    return tokens


# ============================================================
# Document element iteration
# ============================================================

def iter_block_items(doc):
    """Iterate paragraphs and tables in document order.

    python-docx's doc.paragraphs and doc.tables don't interleave,
    so we walk the XML body children directly.
    """
    from docx.table import Table as DocxTable
    from docx.text.paragraph import Paragraph as DocxParagraph

    body = doc.element.body
    for child in body:
        if child.tag == qn("w:p"):
            yield DocxParagraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield DocxTable(child, doc)


def is_paragraph(elem) -> bool:
    """Check if an element is a Paragraph (not a Table)."""
    return hasattr(elem, "runs")


def is_table(elem) -> bool:
    """Check if an element is a Table (not a Paragraph)."""
    return hasattr(elem, "rows") and not hasattr(elem, "runs")


# ============================================================
# Main parser
# ============================================================

def parse_docx(path: str) -> ParseResult:
    """Parse a SNIFFer DOCX file into structured data."""
    doc = Document(path)
    errors: list[str] = []
    warnings: list[str] = []

    # Pre-flight checks
    if has_track_changes(doc):
        errors.append("Document contains tracked changes. Accept or reject all changes before parsing.")
    if has_comments(doc):
        errors.append("Document contains comments. Resolve or remove all comments before parsing.")

    # Walk the document
    elements = list(iter_block_items(doc))
    n = len(elements)
    i = 0

    legend = Legend()
    overrides: list[LegendOverride] = []
    sections: list[BandSection] = []
    found_global_legend = False
    global_legend_first = True
    pending_subject: Optional[str] = None  # subject line waiting to attach to next band section

    # All defined token names (for cross-check)
    defined_tokens: set[str] = set()

    # --- Skip preamble: everything before first EL_LegendHeading ---
    # If any EL_* style appears before EL_LegendHeading, that's a structural error
    while i < n:
        elem = elements[i]
        if is_paragraph(elem):
            style = resolve_el_style(elem)
            if style == "EL_LegendHeading":
                break  # found it, stop skipping
            if style is not None:
                # An EL_* style before LEGEND = ordering violation
                errors.append(
                    f"EL_* element '{style}' (\"{para_text(elem)[:60]}\") appears before "
                    f"the global LEGEND section. LEGEND must be the first EL_* element."
                )
        i += 1

    while i < n:
        elem = elements[i]

        # --- Paragraph ---
        if is_paragraph(elem):
            para = elem
            style = resolve_el_style(para)
            text = para_text(para)

            # Skip empty non-EL paragraphs (spacers)
            if not style and is_empty_para(para):
                i += 1
                continue

            # EL_Rule — consumed during table detection, not emitted
            if style == "EL_Rule":
                # Check for table sandwich: EL_Rule → table → EL_Rule
                if i + 2 < n and is_table(elements[i + 1]):
                    if is_paragraph(elements[i + 2]):
                        closing_style = resolve_el_style(elements[i + 2])
                        if closing_style != "EL_Rule":
                            errors.append(
                                f"Table at position {i+1} has opening EL_Rule but closing "
                                f"paragraph is '{closing_style or get_style_name(elements[i+2])}', expected EL_Rule"
                            )
                    i += 1  # skip to table, handled below
                    continue
                else:
                    errors.append(f"Orphaned EL_Rule at position {i} — no adjacent table")
                    i += 1
                    continue

            # EL_LegendHeading
            if style == "EL_LegendHeading":
                m = LEGEND_HEADING_RE.match(text)
                if not m:
                    errors.append(f"EL_LegendHeading text '{text}' doesn't match pattern 'LEGEND' or 'LEGEND - <Language>'")
                    i += 1
                    continue

                lang = m.group(1)
                if lang is None:
                    # Global legend
                    if found_global_legend:
                        errors.append("Multiple global LEGEND sections found")
                    found_global_legend = True
                    if sections or overrides:
                        global_legend_first = False

                    # Consume sub-tables
                    i += 1
                    legend, i = _parse_legend_tables(elements, i, n, errors, "Global LEGEND")
                    for dt in legend.dynamic_tokens:
                        defined_tokens.add(dt.token)
                    for lt in legend.link_tokens:
                        defined_tokens.add(lt.token)
                    continue
                else:
                    # Per-language override
                    lang = lang.strip()
                    if lang not in VALID_LANGUAGES:
                        errors.append(f"LEGEND override for unknown language '{lang}'")

                    i += 1
                    override_legend, i = _parse_legend_tables(
                        elements, i, n, errors, f"LEGEND - {lang}"
                    )
                    override = LegendOverride(
                        language=lang,
                        dynamic_tokens=override_legend.dynamic_tokens,
                        link_tokens=override_legend.link_tokens
                    )
                    overrides.append(override)
                    continue

            # EL_SubjectHeading
            if style == "EL_SubjectHeading":
                m = SUBJECT_HEADING_RE.match(text)
                if not m:
                    errors.append(f"EL_SubjectHeading text '{text}' doesn't match pattern 'SUBJECT'")
                    i += 1
                    continue

                i += 1
                # Next paragraph should be EL_Subject
                if i < n and is_paragraph(elements[i]):
                    subj_style = resolve_el_style(elements[i])
                    if subj_style == "EL_Subject":
                        pending_subject = para_text(elements[i])
                        i += 1
                    else:
                        errors.append(f"Expected EL_Subject after EL_SubjectHeading, got '{subj_style or get_style_name(elements[i])}'")
                else:
                    errors.append("EL_SubjectHeading at end of document with no EL_Subject paragraph")
                continue

            # EL_Subject outside of SubjectHeading context
            if style == "EL_Subject":
                errors.append(f"EL_Subject paragraph without preceding EL_SubjectHeading: \"{text[:60]}\"")
                i += 1
                continue

            # EL_BandHeading
            if style == "EL_BandHeading":
                m = BAND_HEADING_RE.match(text)
                if not m:
                    errors.append(f"EL_BandHeading text '{text}' doesn't match pattern 'Band N - <Language>'")
                    i += 1
                    continue

                band_num = int(m.group(1))
                lang = m.group(2).strip()
                if lang not in VALID_LANGUAGES:
                    errors.append(f"Band heading references unknown language '{lang}'")

                section = BandSection(band=band_num, language=lang, subject=pending_subject)
                pending_subject = None  # consumed
                i += 1

                # Consume body paragraphs until next heading or end
                while i < n:
                    elem2 = elements[i]
                    if is_table(elem2):
                        # table in body — unexpected, skip with warning
                        warnings.append(f"Unexpected table inside Band {band_num} - {lang}")
                        i += 1
                        continue

                    if not is_paragraph(elem2):
                        i += 1
                        continue

                    s2 = resolve_el_style(elem2)

                    # Next section heading — stop
                    if s2 in SECTION_HEADING_STYLES:
                        break

                    # EL_Rule in body — shouldn't be here
                    if s2 == "EL_Rule":
                        errors.append(f"EL_Rule inside Band {band_num} - {lang} body (position {i})")
                        i += 1
                        continue

                    # Skip empty paragraphs (spacers)
                    if is_empty_para(elem2) and s2 is None:
                        i += 1
                        continue

                    # Non-EL style on a structural element
                    if s2 is None and not is_empty_para(elem2):
                        preview = para_text(elem2)[:60]
                        errors.append(
                            f"Non-EL_* style '{get_style_name(elem2) or 'Normal'}' on paragraph "
                            f"in Band {band_num} - {lang}: \"{preview}\""
                        )
                        i += 1
                        continue

                    # Valid body paragraph
                    if s2 in BAND_STYLES:
                        runs = process_runs(elem2)
                        if runs:  # skip if empty after processing
                            section.paragraphs.append(
                                ParsedParagraph(style=s2, runs=runs)
                            )
                        i += 1
                        continue

                    # EL_LegendLabel inside body — wrong place
                    if s2 == "EL_LegendLabel":
                        errors.append(f"EL_LegendLabel inside Band {band_num} - {lang} body")
                        i += 1
                        continue

                    i += 1

                if not section.paragraphs:
                    warnings.append(f"Band {band_num} - {lang} has no body content")

                sections.append(section)
                continue

            # Non-EL paragraph with content — error
            if not style and not is_empty_para(para):
                preview = text[:60]
                errors.append(
                    f"Non-EL_* style '{get_style_name(para) or 'Normal'}' on paragraph: \"{preview}\""
                )
                i += 1
                continue

            # EL_LegendLabel outside legend context — error
            if style == "EL_LegendLabel":
                errors.append(f"EL_LegendLabel outside of a LEGEND section: \"{text[:60]}\"")
                i += 1
                continue

            i += 1
            continue

        # --- Table outside legend context ---
        # Tables should only appear inside the EL_Rule sandwich.
        # If we get here, it's a table not preceded by EL_Rule.
        if is_table(elem):
            errors.append(f"Table at position {i} not bounded by EL_Rule sandwich")
            i += 1
            continue

        i += 1

    # Post-parse validation
    if not found_global_legend:
        errors.append("Document is missing the global LEGEND section")
    elif not global_legend_first:
        errors.append("Global LEGEND section is not the first section in the document")

    # Cross-check: override tokens must exist in global legend
    for ov in overrides:
        for dt in ov.dynamic_tokens:
            if dt.token not in defined_tokens:
                errors.append(
                    f"LEGEND - {ov.language} defines token '{dt.token}' not in global LEGEND"
                )
        for lt in ov.link_tokens:
            if lt.token not in defined_tokens:
                errors.append(
                    f"LEGEND - {ov.language} defines token '{lt.token}' not in global LEGEND"
                )

    # Cross-check: tokens used in body must be in global legend
    used_tokens: set[str] = set()
    for section in sections:
        for para in section.paragraphs:
            for run in para.runs:
                if run.token:
                    used_tokens.add(run.text)
                    if run.text not in defined_tokens:
                        errors.append(
                            f"Token '{run.text}' used in Band {section.band} - "
                            f"{section.language} not defined in global LEGEND"
                        )

    # Soft: tokens defined but never used
    unused = defined_tokens - used_tokens
    for tok in sorted(unused):
        warnings.append(f"Token '{tok}' defined in global LEGEND but never used in any band section")

    # Soft: override rows identical to global
    global_dynamic = {dt.token: dt for dt in legend.dynamic_tokens}
    global_link = {lt.token: lt for lt in legend.link_tokens}
    for ov in overrides:
        for dt in ov.dynamic_tokens:
            g = global_dynamic.get(dt.token)
            if g and dt.example_value == g.example_value and dt.description == g.description:
                warnings.append(
                    f"LEGEND - {ov.language}: override for '{dt.token}' is identical to global (no-op)"
                )
        for lt in ov.link_tokens:
            g = global_link.get(lt.token)
            if (g and lt.display_text == g.display_text
                    and lt.target_url == g.target_url
                    and lt.description == g.description):
                warnings.append(
                    f"LEGEND - {ov.language}: override for '{lt.token}' is identical to global (no-op)"
                )

    return ParseResult(
        legend=legend,
        overrides=overrides,
        sections=sections,
        hard_errors=errors,
        soft_warnings=warnings
    )


def _parse_legend_tables(elements, i, n, errors, context) -> tuple[Legend, int]:
    """Parse legend sub-tables starting at position i.

    Expects: EL_LegendLabel → EL_Rule → table → EL_Rule, repeated for each sub-table.
    Returns the Legend and the new position.
    """
    legend = Legend()
    while i < n:
        elem = elements[i]

        # Stop at next section heading
        if is_paragraph(elem):
            style = resolve_el_style(elem)
            if style in SECTION_HEADING_STYLES:
                break

            if style == "EL_LegendLabel":
                label_text = para_text(elem).lower()
                label_display = para_text(elem)
                i += 1

                # Expect EL_Rule
                if i >= n:
                    errors.append(f"{context}: EL_LegendLabel '{label_display}' at end of document with no table")
                    break
                rule_elem = elements[i]
                if is_paragraph(rule_elem) and resolve_el_style(rule_elem) == "EL_Rule":
                    i += 1
                else:
                    errors.append(f"{context}: Expected EL_Rule after EL_LegendLabel, got '{get_style_name(rule_elem) if is_paragraph(rule_elem) else 'table'}'")

                # Expect table
                if i >= n or not is_table(elements[i]):
                    errors.append(f"{context}: Expected table after EL_Rule for '{label_display}'")
                    continue
                table = elements[i]
                rows = read_table_rows(table)
                i += 1

                # Expect closing EL_Rule
                if i < n and is_paragraph(elements[i]) and resolve_el_style(elements[i]) == "EL_Rule":
                    i += 1
                else:
                    errors.append(f"{context}: Missing closing EL_Rule after table for '{label_display}'")

                # Dispatch on label
                if "dynamic" in label_text:
                    legend.dynamic_tokens = parse_dynamic_token_table(rows, errors, context)
                elif "link" in label_text:
                    legend.link_tokens = parse_link_token_table(rows, errors, context)
                else:
                    errors.append(f"{context}: Unrecognized legend label '{label_display}' — expected 'Dynamic Data Tokens' or 'Link Tokens'")

                continue

            # Skip empty paragraphs / spacers
            if is_empty_para(elem):
                i += 1
                continue

            # Skip EL_Rule (might be standalone — error handled elsewhere)
            if style == "EL_Rule":
                i += 1
                continue

            # Anything else — we've left the legend section
            break

        # Table without EL_Rule lead-in (shouldn't happen in well-formed doc)
        if is_table(elem):
            errors.append(f"{context}: Table without EL_Rule sandwich at position {i}")
            i += 1
            continue

        i += 1

    return legend, i


# ============================================================
# Serialization
# ============================================================

def to_dict(result: ParseResult) -> dict:
    """Convert ParseResult to a JSON-serializable dict."""
    def _run(r: Run) -> dict:
        d = {"text": r.text, "bold": r.bold, "italic": r.italic}
        if r.token:
            d["token"] = True
        return d

    def _para(p: ParsedParagraph) -> dict:
        return {"style": p.style, "runs": [_run(r) for r in p.runs]}

    def _dt(d: DynamicToken) -> dict:
        return {"token": d.token, "example_value": d.example_value, "description": d.description}

    def _lt(l: LinkToken) -> dict:
        return {"token": l.token, "display_text": l.display_text,
                "target_url": l.target_url, "description": l.description}

    legend_dict = {
        "dynamic_tokens": [_dt(d) for d in result.legend.dynamic_tokens],
        "link_tokens": [_lt(l) for l in result.legend.link_tokens],
    }

    overrides_dict = {}
    for ov in result.overrides:
        ov_entry = {}
        if ov.dynamic_tokens:
            ov_entry["dynamic_tokens"] = [_dt(d) for d in ov.dynamic_tokens]
        if ov.link_tokens:
            ov_entry["link_tokens"] = [_lt(l) for l in ov.link_tokens]
        overrides_dict[ov.language] = ov_entry

    if overrides_dict:
        legend_dict["overrides"] = overrides_dict

    sections_list = []
    for s in result.sections:
        entry = {
            "band": s.band,
            "language": s.language,
            "paragraphs": [_para(p) for p in s.paragraphs]
        }
        if s.subject is not None:
            entry["subject"] = s.subject
        sections_list.append(entry)

    return {
        "legend": legend_dict,
        "sections": sections_list,
    }


# ============================================================
# CLI
# ============================================================

def main():
    parser = argparse.ArgumentParser(description="Parse a SNIFFer DOCX into structured runs")
    parser.add_argument("docx_path", help="Path to the .docx file")
    parser.add_argument("--output", "-o", help="Output JSON path (default: stdout)")
    args = parser.parse_args()

    path = Path(args.docx_path)
    if not path.exists():
        print(f"Error: file not found: {path}", file=sys.stderr)
        sys.exit(4)

    result = parse_docx(str(path))

    # Report errors
    if result.hard_errors:
        print(f"\n{'='*60}", file=sys.stderr)
        print(f"HARD ERRORS ({len(result.hard_errors)}) — document rejected", file=sys.stderr)
        print(f"{'='*60}", file=sys.stderr)
        for err in result.hard_errors:
            print(f"  ✗ {err}", file=sys.stderr)

    if result.soft_warnings:
        print(f"\n{'='*60}", file=sys.stderr)
        print(f"SOFT WARNINGS ({len(result.soft_warnings)})", file=sys.stderr)
        print(f"{'='*60}", file=sys.stderr)
        for warn in result.soft_warnings:
            print(f"  ⚠ {warn}", file=sys.stderr)

    if result.hard_errors:
        print(f"\nParse failed with {len(result.hard_errors)} error(s). No output emitted.",
              file=sys.stderr)
        sys.exit(1)

    # Emit output
    output = to_dict(result)
    json_str = json.dumps(output, indent=2, ensure_ascii=False)

    if args.output:
        out_path = Path(args.output)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_text(json_str, encoding="utf-8")
        print(f"Written to {out_path}", file=sys.stderr)
    else:
        print(json_str)

    sys.exit(2 if result.soft_warnings else 0)


if __name__ == "__main__":
    main()
